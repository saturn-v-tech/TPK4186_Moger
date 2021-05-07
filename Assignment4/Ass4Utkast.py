# Assignment 4 - TPK4186 - Assignment 4
# Torje Furu StudentID: 768765
# Lars Gruben StudentID: 473811

# 1. Imported packages
#----------------------

import sys
import re
import matplotlib.pyplot as plt
import numpy as np
from pathlib import Path
from sklearn import linear_model, svm, neural_network, gaussian_process
# from sklearn.naive_bayes import GaussianNB
import statistics
from sklearn.svm import NuSVR
from sklearn.linear_model import SGDClassifier, LogisticRegression
from sklearn.neural_network import MLPRegressor
from sklearn.gaussian_process import GaussianProcessClassifier
from sklearn.linear_model import LinearRegression
from sklearn.tree import DecisionTreeRegressor, DecisionTreeClassifier
from docx import Document

# Parser
#--------

class Parser:

  def __init__(self):
    self.expectedDuration = None
    self.dataList = None
    self.project = None
    self.delay = None
    self.durationList = []
    self.expectedDurations = []
    self.progressList = []
    self.delays = []
    self.projectList = []
    self.duration = None
    self.trainingList = []
    self.lowerLimitTraining = 0.2
    self.upperLimitTraining = 0.25
    self.lowerLimitTest = 0.2
    self.upperLimitTest = 0.25
    self.count = 0
    self.count1 = 0

  def ImportFile(self, inputFileName, outputFileName, outputFileName2, outputFileName3, outputFileName4):
    try:
      inputFile = open(inputFileName, 'r')
    except:
      sys.stderr.write('unable to open file ' + inputFileName + '\n')
      sys.exit()
    project = self.ReadFile(inputFile, outputFileName, outputFileName2, outputFileName3, outputFileName4)
    inputFile.flush()
    inputFile.close()
    return project

  def OpenFiles(self, directory, outputFilename, outputFileName2, outputFileName3, outputFileName4):
    for path in Path(directory).iterdir():
      self.ImportFile(path, outputFileName, outputFileName2, outputFileName3, outputFileName4)


  def ReadFile(self, inputFile, outputFile, outputFile2, outpuFile3, outputFile4):
    state = 0
    week = []
    foundation = []
    framing = []
    curtainWall = []
    HVAC = []
    fireFighting = []
    elevator = []
    electrical = []
    ArchitecturalFinishing = []
    lists = [week, foundation, framing, curtainWall, HVAC, fireFighting, elevator, electrical, ArchitecturalFinishing]
    for line in inputFile:
      line = line.rstrip()
      if re.match('Name', line):
        line = line.split()
        self.project = line[1]
      elif re.match('Expected duration', line):
        line = line.split()
        self.expectedDuration = int(line[2])
        state = 1
      elif state == 1:
        line = line.split()
        for i in range(len(line)):
          lists[i].append(line[i])
    i = list(range(0,len(lists)))
    self.week = lists[0]
    self.foundation =lists[1]
    zippedList = list(zip(lists[i[0]],lists[i[1]], lists[i[2]], lists[i[3]], lists[i[4]], lists[i[5]], lists[i[6]], lists[i[7]], lists[i[8]]))
    self.projectList.append(self.project)
    self.dataList = zippedList
    length = len(self.GetDataList()) - 1
    self.delay = int(self.dataList[length][0]) - int(self.expectedDuration)
    self.delays.append(self.delay)
    self.duration = int(self.dataList[length][0])
    statistic.PercentageOfError(self.duration, int(self.expectedDuration))
    self.expectedDurations.append(self.expectedDuration)
    self.PsvFileDataToPredictFailure(outputFile, outputFile2)
    self.PsvFileDataToPredictDuration(outpuFile3, outputFile4)


  def PsvFileDataToPredictDuration(self, outputFileName, outputFileName2):
    try:
     outputFile = open(outputFileName, 'a')
     outputFile2 = open(outputFileName2, 'a')
    except:
      sys.stderr.write('unable to open file ' + outputFile + '\n')
      sys.stderr.write('unable to open file ' + outputFile2 + '\n')
      sys.exit()
    project = self.TrainingAndTestDataToCsvDuration(outputFile, outputFile2)
    outputFile.flush()
    outputFile.close()
    outputFile2.flush()
    outputFile2.close()
    return project

  def PsvFileDataToPredictFailure(self, outputFileName, outputFileName2):
    try:
     outputFile = open(outputFileName, 'a')
     outputFile2 = open(outputFileName2, 'a')
    except:
      sys.stderr.write('unable to open file ' + outputFile + '\n')
      sys.stderr.write('unable to open file ' + outputFile2 + '\n')
      sys.exit()
    project = self.TrainingAndTestDataToCsvFailure(outputFile, outputFile2)
    outputFile.flush()
    outputFile.close()
    outputFile2.flush()
    outputFile2.close()
    return project

  def TrainingAndTestDataToCsvFailure(self, outputFile, outputFile2):
    self.count +=1
    for i in range(1,len(self.dataList)):
      if self.count >=100:
        if self.lowerLimitTest < int(self.dataList[i][0])/int(self.expectedDuration) <  self.upperLimitTest:
          normalizedValue = round(int(self.dataList[i][0])/int(self.expectedDuration),2)
          outputFile2.write(f"{normalizedValue},{self.dataList[i][1]}, {self.dataList[i][2]}, {self.dataList[i][3]}, {self.dataList[i][4]}, {self.dataList[i][5]}, {self.dataList[i][6]}, {self.dataList[i][7]}, {self.dataList[i][8]}")
          if self.duration/int(self.expectedDuration) < 1.4:
            outputFile2.write(",1\n")
          elif self.duration/int(self.expectedDuration) >= 1.4:
            outputFile2.write(",0 \n")
      else:
        if self.lowerLimitTraining < int(self.dataList[i][0])/int(self.expectedDuration) < self.upperLimitTraining:
          normalizedValue = round(int(self.dataList[i][0])/int(self.expectedDuration),2)
          outputFile.write(f"{normalizedValue},{self.dataList[i][1]}, {self.dataList[i][2]}, {self.dataList[i][3]}, {self.dataList[i][4]}, {self.dataList[i][5]}, {self.dataList[i][6]}, {self.dataList[i][7]}, {self.dataList[i][8]}")
          if self.duration/int(self.expectedDuration) < 1.4:
            outputFile.write(",1\n")
          elif self.duration/int(self.expectedDuration) >= 1.4:
            outputFile.write(",0 \n")

  def TrainingAndTestDataToCsvDuration(self, outputFile, outputFile2):
    self.count1 +=1
    for i in range(1,len(self.dataList)):
      if self.count >=100:
        if self.lowerLimitTest < int(self.dataList[i][0])/int(self.expectedDuration) < self.upperLimitTest:
          normalizedValue = round(int(self.dataList[i][0])/int(self.expectedDuration),2)
          outputFile2.write(f"{normalizedValue},{self.dataList[i][1]}, {self.dataList[i][2]}, {self.dataList[i][3]}, {self.dataList[i][4]}, {self.dataList[i][5]}, {self.dataList[i][6]}, {self.dataList[i][7]}, {self.dataList[i][8]}, {int(self.duration)}\n")
      else:
        if self.lowerLimitTraining < int(self.dataList[i][0])/int(self.expectedDuration) < self.upperLimitTraining:
          normalizedValue = round(int(self.dataList[i][0])/int(self.expectedDuration),2)
          outputFile.write(f"{normalizedValue},{self.dataList[i][1]}, {self.dataList[i][2]}, {self.dataList[i][3]}, {self.dataList[i][4]}, {self.dataList[i][5]}, {self.dataList[i][6]}, {self.dataList[i][7]}, {self.dataList[i][8]}, {int(self.duration)}\n")

  def GetTrainingList(self):
    return self.trainingList

  def GetDurationList(self):
    return self.durationList

  def GetDataList(self):
    return self.dataList

  def GetExpectedDuration(self):
    return self.expectedDuration

  def GetProject(self):
    return self.project

  def GetDelay(self):
    return self.delay

  def GetListOfDelays(self):
    return self.delays

class Statistic:

  def __init__(self):
    self.percentageOfErrors = []
    self.plotFileNameDelay = None
    self.plotFileNamePercantage = None

  def PercentageOfError(self, duration, expectedDuration):
    Error = duration/expectedDuration
    percentageOfError = round((Error - 1) * 100,2)
    return self.percentageOfErrors.append(percentageOfError)

  def GetListOfPercentageOfError(self):
    return self.percentageOfErrors

  def PlotWeeksOfDelay(self):
    plt.hist(parser.GetListOfDelays(), bins = 100)
    plt.title('Project delay')
    plt.xlabel('Week of delays')
    plt.ylabel('Number of projects')
    plt.xticks(np.arange(0, max(parser.GetListOfDelays()), 10))
    plt.savefig('projectWeekDelay.jpeg')
    fileName = "projectWeekDelay.jpeg"
    self.plotFileNameDelay = fileName
    plt.show()

  def PlotPercentageOfDelay(self):
    plt.hist(self.GetListOfPercentageOfError(), bins = 100)
    plt.title('Project delay')
    plt.xlabel('Percentage')
    plt.ylabel('Number of projects')
    plt.xticks(np.arange(0,max(self.GetListOfPercentageOfError()),5))
    plt.savefig('projectPercentageDelay.jpeg')
    fileName = "projectPercentageDelay.jpeg"
    self.plotFileNamePercantage = fileName
    plt.show()

class Printer:


  def PrintResults(self, actualResult, predictedLabels):
    predictionsWithSuccess = 0
    predictionsWithfailure = 0
    for i in range(0, len(actualResult)):
      if actualResult[i] == predictedLabels[i]:
        predictionsWithSuccess += 1
      else:
        predictionsWithfailure += 1
    sys.stdout.write(f"{predictionsWithSuccess} simulations was successed\n")
    sys.stdout.write(f"{predictionsWithfailure} simulations was failed\n")

  def PrintProsentageOfFail(self, predictedRewards, actualDuration):
    prosentageOfActualTime = []
    predictionWithSuccess = 0
    predictionsWithfailure = 0
    for i in range(0, len(actualDuration)):
      if 0.9 < round(predictedRewards[i]/actualDuration[i],3) < 1.1:
        predictionWithSuccess += 1
      else:
        predictionsWithfailure += 1
    sys.stdout.write(f"{predictionWithSuccess} was predicted with less than 10% failure\n")
    sys.stdout.write(f"{predictionsWithfailure} was predicted with more than 10% failure\n")

class DocumentGenerator:

  def __init__(self):
    self.countClass = 0
    self.countReg = 0

  def GenerateReport(self, fileName, testNamesClass, testNamesReg, actualResult, actualDuration, predictedLabelsList, predictedRewardsList):
    document = Document()
    document.add_heading('Experiments on durationtime of projects')
    document.add_paragraph(f"The training simulations is performed with lower limit {parser.lowerLimitTraining} and upper limit {parser.upperLimitTraining}")
    document.add_paragraph(f"The test simulations is performed with lower limit {parser.lowerLimitTest} and upper limit {parser.upperLimitTest}")
    for simulationClassificationName in testNamesClass:
      if self.countClass == 0:
        document.add_heading("Simulations of failed projects", level = 1)
      document.add_heading(f"{simulationClassificationName}", level = 2)
      self.PrintClassification(actualResult, predictedLabelsList, document)
      self.countClass += 1
    for simulationRegressionName in testNamesReg:
      if self.countReg == 0:
        document.add_heading("Simulations of predicted durationtime", level = 1)
      document.add_heading(f"{simulationRegressionName}", level = 2)
      self.PrintRegression(actualDuration, predictedRewardsList, document)
      self.countReg += 1
    document.add_heading("Graph with the exceeded percentage of expected duration")
    document.add_picture(statistic.plotFileNamePercantage)
    document.add_heading("Graph with the number of weeks exceed of the expected duration time")
    document.add_picture(statistic.plotFileNameDelay)
    document.save(fileName)


  def PrintClassification(self, actualResult, predictedLabelsList, document):
    predictionsWithSuccess = 0
    predictionsWithfailure = 0
    predictedLabels = predictedLabelsList[self.countClass][0]
    for i in range(0, len(actualResult)):
      if actualResult[i] == predictedLabels[i]:
        predictionsWithSuccess += 1
      else:
        predictionsWithfailure += 1
    document.add_paragraph(f"{predictionsWithSuccess} simulations was successed")
    document.add_paragraph(f"{predictionsWithfailure} simulations was failed")

  def PrintRegression(self, actualDuration, predictedRewardsList, document):
    predictionWithSuccess = 0
    predictionsWithfailure = 0
    predictedRewards = predictedRewardsList[self.countReg][0]
    for i in range(0, len(actualDuration)):
      if 0.9 < round(predictedRewards[i]/actualDuration[i],3) < 1.1:
        predictionWithSuccess += 1
      else:
        predictionsWithfailure += 1
    document.add_paragraph(f"{predictionWithSuccess} was predicted with less than 10% failure")
    document.add_paragraph(f"{predictionsWithfailure} was predicted with more than 10% failure")


# Defining classes
# ------------------
parser = Parser()
statistic = Statistic()
documentGenerator = DocumentGenerator()
printer = Printer()

# Documentnames
# -------------

inputFileName = 'Project Data'
outputFileName = 'trainingDataForFailures.csv'
outputFileName2 = 'testDataForFailures.csv'
outputFileName3 = 'trainingDataForDuration.csv'
outputFileName4 = 'testDataForDuration.csv'
documentName = "Assignment 4 - Torje and Lars - TPK4186.docx"


# Creating training and testdata
# ------------------------------
parser.OpenFiles(inputFileName, outputFileName, outputFileName2, outputFileName3, outputFileName4)

# Creating the graph
# --------------------
statistic.PlotPercentageOfDelay()
statistic.PlotWeeksOfDelay()
#  Creating lists to store results
# ------------------
predictedRewardsList = [[] for i in range(6)]
predictedLabelsList = [[]for i in range(4)]

# MachineLearning
# ----------------
# Trainingdata - Classification
# ------------------------------
trainingData = np.genfromtxt(outputFileName, delimiter = ",")
trainingData = np.array(trainingData)
trainingInstances = trainingData[:, 0:-2]
trainingLables = trainingData[:, -1]

# Testdata
#----------
testData = np.genfromtxt(outputFileName2, delimiter = ",")
testData = np.array(testData)
testInstances = testData[:, 0:-2]
actualResult = testData[:, -1]

# First Test - Linear classificaton
# --------------------------------

testName1class ="First Test - Linear classificaton"
model = SGDClassifier(loss='hinge', penalty='l2', alpha=0.0001, l1_ratio=0.15, fit_intercept=True, max_iter=1000, tol=0.001, shuffle=True, verbose=0, epsilon=0.1, n_jobs=None, random_state=None, learning_rate='optimal', eta0=0.0, power_t=0.5, early_stopping=False, validation_fraction=0.1, n_iter_no_change=5, class_weight=None, warm_start=False, average=False)
# model = linear_model.LogisticRegression()
model.fit(trainingInstances, trainingLables)
predictedLabels = model.predict(testInstances)
predictedLabelsList[0].append(predictedLabels)
# printer.PrintResults(actualResult, predictedLabels)

# Second Test - Support Vector machines
# -------------------------------------

testName2class = "Second Test - Support Vector machines"
model = svm.SVC()
model.fit(trainingInstances, trainingLables)
predictedLabels = model.predict(testInstances)
predictedLabelsList[1].append(predictedLabels)
# printer.PrintResults(actualResult, predictedLabels)

# Third Test - Gaussian Process
# -------------------------------------
testName3class = "Third Test - Gaussian Process"
model = GaussianProcessClassifier(kernel=None, optimizer='fmin_l_bfgs_b', n_restarts_optimizer=0, max_iter_predict=1000, warm_start=False, copy_X_train=True, random_state=None, multi_class='one_vs_rest', n_jobs=None)
model.fit(trainingInstances, trainingLables)
predictedLabels = model.predict(testInstances)
predictedLabelsList[2].append(predictedLabels)
# printer.PrintResults(actualResult, predictedLabels)

# Fourth Test - Decision Tree
# -----------------------------
testName4class = "Fourth Test - Decision Tree"

model = DecisionTreeClassifier( criterion='gini', splitter='best', max_depth=None, min_samples_split=2, min_samples_leaf=1, min_weight_fraction_leaf=0.0, max_features=None, random_state=None, max_leaf_nodes=None, min_impurity_decrease=0.0, min_impurity_split=None, class_weight=None, ccp_alpha=0.0)
model.fit(trainingInstances, trainingLables)
predictedLabels = model.predict(testInstances)
predictedLabelsList[3].append(predictedLabels)
# printer.PrintResults(actualResult, predictedLabels)

# Regression
# ------------

# Trainingdata - Regression
# --------------------------

trainingData = np.genfromtxt(outputFileName3, delimiter = ",")
trainingData = np.array(trainingData)
trainingInstances = trainingData[:, 0:-2]
trainingLables = trainingData[:, -1]


# Testdata
# ----------
testData = np.genfromtxt(outputFileName4, delimiter = ",")
testData = np.array(testData)
testInstances = testData[:, 0:-2]
actualDuration = testData[:, -1]

# First Test - Linear Regression
# -------------------------------
testName1Reg = "First Test - Linear Regression"

regr = LinearRegression(fit_intercept=True, normalize=False, copy_X=True, n_jobs=None)
regr.fit(trainingInstances,trainingLables)

predictedRewards = regr.predict(testInstances)
predictedRewardsList[0].append(predictedRewards)
# printer.PrintProsentageOfFail(predictedRewards, actualDuration)

# Second Test - Nu Support Vector Regression
# ---------------------------------------------

testName2Reg = "Second Test - Nu Support Vector Regression "
regr = NuSVR(nu=0.5, C=1.0, kernel='rbf', degree=3, gamma='scale', coef0=0.0, shrinking=True, tol=0.001, cache_size=200, verbose=False, max_iter=1000)
regr.fit(trainingInstances, trainingLables)

predictedRewards = regr.predict(testInstances)
predictedRewardsList[1].append(predictedRewards)
# printer.PrintProsentageOfFail(predictedRewards, actualDuration)

# Third Test - Neutral Network
# ------------------------------
testName3Reg = "Third Test - Neutral Network"
regr = MLPRegressor(hidden_layer_sizes=100, activation='relu', solver='adam', alpha=0.0001, batch_size='auto', learning_rate='constant', learning_rate_init=0.001, power_t=0.5, max_iter=1000, shuffle=True, random_state=None, tol=0.0001, verbose=False, warm_start=False, momentum=0.9, nesterovs_momentum=True, early_stopping=False, validation_fraction=0.1, beta_1=0.9, beta_2=0.999, epsilon=1e-08, n_iter_no_change=10, max_fun=15000)
regr.fit(trainingInstances,trainingLables)

predictedRewards = regr.predict(testInstances)
predictedRewardsList[2].append(predictedRewards)
# printer.PrintProsentageOfFail(predictedRewards, actualDuration)

# Fourth Test - Passive Aggressive
# ------------------------------
testName4Reg = "Fourth Test - Passive Aggressive"
regr = linear_model.PassiveAggressiveRegressor(C=1.0, fit_intercept=True, max_iter=1000, tol=0.002, early_stopping=False, validation_fraction=0.1, n_iter_no_change=100, shuffle=True, verbose=0, loss='epsilon_insensitive', epsilon=0.1, random_state=None, warm_start=False, average=False)
regr.fit(trainingInstances,trainingLables)

predictedRewards = regr.predict(testInstances)
predictedRewardsList[3].append(predictedRewards)
# printer.PrintProsentageOfFail(predictedRewards, actualDuration)

# Fifth Test - DecisionTree
# --------------------------
testName5Reg = "Fifth Test - DecisionTree"
regr = DecisionTreeRegressor(criterion='mse', splitter='best', max_depth=None, min_samples_split=2, min_samples_leaf=1, min_weight_fraction_leaf=0.0, max_features=None, random_state=None, max_leaf_nodes=None, min_impurity_decrease=0.0, min_impurity_split=None, ccp_alpha=0.0)
regr.fit(trainingInstances,trainingLables)

predictedRewards = regr.predict(testInstances)
predictedRewardsList[4].append(predictedRewards)
# printer.PrintProsentageOfFail(predictedRewards, actualDuration)

# Sixth Test - Logistic Regression
# -------------------------
testName6Reg = "Sixth Test - Logistic Regression"
regr = LogisticRegression(penalty='l2', dual=False, tol=0.0001, C=1.0, fit_intercept=True, intercept_scaling=1, class_weight=None, random_state=None, solver='lbfgs', max_iter=5000, multi_class='auto', verbose=0, warm_start=False, n_jobs=None, l1_ratio=None)
regr.fit(trainingInstances,trainingLables)

predictedRewards = regr.predict(testInstances)
predictedRewardsList[5].append(predictedRewards)
# printer.PrintProsentageOfFail(predictedRewards, actualDuration)





# Saving values to use in making the document
# -------------------------------------------

testNamesClass = [testName1class, testName2class, testName3class, testName4class]
testNamesReg = [testName1Reg, testName2Reg, testName3Reg, testName4Reg, testName5Reg, testName6Reg]
documentGenerator.GenerateReport(documentName, testNamesClass, testNamesReg, actualResult, actualDuration, predictedLabelsList, predictedRewardsList)